import os
from pathlib import Path
import docx
from PyPDF2 import PdfReader
from typing import Any, List, Dict, Union, TextIO, BinaryIO

from langchain.chat_models import init_chat_model
from langchain_core.messages import HumanMessage
from pydantic import BaseModel, Field
from enums import Proficiency
from datetime import datetime
import json


class CleanResponse(BaseModel):
    job_role: str = Field(description="The job role extracted from the content")
    cleaned_description: str = Field(
        description="Clean job description with only relevant information"
    )


class JobDetailsResponse(BaseModel):
    job_role: str = Field(description="The job role extracted from the content")
    company_name: str = Field(description="The company name extracted from the content")
    job_location: str = Field(description="The job location extracted from the content")
    description_summary: str = Field(description="Short job description summary")


class Skill(BaseModel):
    skill_name: str = Field(description="The name of the skill")
    proficiency_level: str = Field(
        description="Required proficiency: beginner, intermediate, advanced, or expert"
    )


class Requirements(BaseModel):
    skills: List[Skill] = Field(description="List of skills required for the job")


class SkillGap(BaseModel):
    skill_name: str = Field(description="The name of the skill with a gap")
    current_level: str = Field(
        description="User's current level: beginner, intermediate, advanced, or expert"
    )
    required_level: str = Field(
        description="Required proficiency level for the job: beginner, intermediate, advanced, or expert"
    )


class Gaps(BaseModel):
    skills: List[SkillGap] = Field(
        description="List of skill gaps between user's current skills and job requirements"
    )


class UserSkill(BaseModel):
    skill_name: str = Field(description="The name of the skill from the CV")
    proficiency_level: str = Field(
        description="User's proficiency in this skill: beginner, intermediate, advanced, or expert"
    )


class CVProfile(BaseModel):
    user_name: str = Field(description="The full name of the user as found in the CV")
    skills: List[UserSkill] = Field(
        description="List of skills and user's proficiency levels"
    )
    profile_summary: str = Field(
        description="A short 3–4 sentence summary of the user's profile, work experience, and strengths"
    )


class Analyzer:
    def __init__(self, llm_api_key):
        self.llm = init_chat_model(
            model="gemini-2.5-flash",
            model_provider="google_genai",
            temperature=0,
            api_key=llm_api_key,
        )

    

    def analyze_cv(self, cv: Union[str, Path, TextIO, BinaryIO]):
        """
        Read a CV file (PDF/DOCX/TXT), extract text, and parse structured user profile.

        Args:
            cv: Path to the CV file (str or Path) OR an open file-like object 
                (TextIO/BinaryIO with a `.name` attribute to infer extension).
        """

        raw_text = None

        if isinstance(cv, (str, Path)):
            file_ext = Path(cv).suffix.lower()

            if file_ext == ".txt":
                with open(cv, "r", encoding="utf-8") as f:
                    raw_text = f.read()

            elif file_ext == ".docx":
                doc = docx.Document(cv)
                raw_text = "\n".join([p.text for p in doc.paragraphs])

            elif file_ext == ".pdf":
                reader = PdfReader(cv)
                raw_text = "\n".join([page.extract_text() or "" for page in reader.pages])

            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

        elif hasattr(cv, "read"):
            name = getattr(cv, "name", None)
            if not name:
                raise ValueError("File-like object must have a 'name' attribute to detect file type")

            file_ext = Path(name).suffix.lower()

            if file_ext == ".txt":
                content = cv.read()
                raw_text = content.decode("utf-8") if isinstance(content, bytes) else content

            elif file_ext == ".docx":
                doc = docx.Document(cv)
                raw_text = "\n".join([p.text for p in doc.paragraphs])

            elif file_ext == ".pdf":
                reader = PdfReader(cv)
                raw_text = "\n".join([page.extract_text() or "" for page in reader.pages])

            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

        else:
            raise TypeError("cv must be a path (str/Path) or a file-like object")

        message = HumanMessage(
            content=f"""
                This is a user's CV/resume. Extract the following structured information:
                - User's full name
                - Skills with proficiency levels (beginner/intermediate/advanced/expert)
                - A short user profile summary of the user.

                CV Content:
                {raw_text}
            """
        )

        structured_llm = self.llm.with_structured_output(CVProfile)
        response = structured_llm.invoke([message])

        return {
            "user_name": response.user_name,
            "skills": [
                {"skill_name": s.skill_name, "proficiency_level": s.proficiency_level}
                for s in response.skills
            ],
            "profile_summary": response.profile_summary,
        }

    def clean_job_description(self, raw_text: str):
        """Use LLM to clean job description text and extract relevant information"""
        message = HumanMessage(
            content=f"""
                This data is from a job posting. Please clean the text and extract only information 
                relevant to the job position and company. Remove navigation elements, advertisements, 
                unrelated content, and formatting artifacts.

                Focus on:
                - Job title/role
                - Job responsibilities and requirements
                - Company information, name, location
                - Skills and qualifications needed
                - Job benefits and conditions

                Raw Text:
                {raw_text}
                """
        )

        structured_llm = self.llm.with_structured_output(CleanResponse)
        response = structured_llm.invoke([message])
        return response.job_role, response.cleaned_description

    def get_job_details(self, job_role, job_description):
        """Use LLM to extract job role, company name, job location, and short job description summary"""
        message = HumanMessage(
            content=f"""
                This is a job description for {job_role} role. Please extract the job role, company name, job location, and short job description summary.

                Description:
                {job_description}
                """
        )

        structured_llm = self.llm.with_structured_output(JobDetailsResponse)
        response = structured_llm.invoke([message])
        return {
            "job_role": response.job_role,
            "company_name": response.company_name,
            "job_location": response.job_location,
            "description_summary": response.description_summary,
        }

    def get_requirements(self, description):
        """Analyze job requirements and extract skills with proficiency levels"""
        message = HumanMessage(
            content=(
                f"Based on the following job description:\n{description}\n\n"
                "Identify the required skills and their required proficiency levels. "
                "For each skill, provide:\n"
                "- skill_name: The name of the skill\n"
                "- proficiency_level: beginner/intermediate/advanced/expert\n"
            )
        )

        structured_llm = self.llm.with_structured_output(Requirements)
        response = structured_llm.invoke([message])

        skills_dicts = []
        for skill in response.skills:
            skill_dict = {
                "skill_name": skill.skill_name.strip().lower(),
                "proficiency_level": skill.proficiency_level.strip().lower(),
            }
            skills_dicts.append(skill_dict)

        return skills_dicts

    def get_gaps(self, user_profile, required_skills):
        """Identify gaps between user's current skills and job requirements"""
        message = HumanMessage(
            content=(
                f"User's current profile: {json.dumps(user_profile['skills'], indent=2)}\n"
                f"Required skills for the job: {json.dumps(required_skills, indent=2)}\n\n"
                "Compare the user's current skills with the job requirements and identify skill gaps. "
                "For each gap, provide:\n"
                "- skill_name: The name of the skill\n"
                "- current_level: User's current proficiency beginner/intermediate/advanced/expert\n"
                "- required_level: Required proficiency level beginner/intermediate/advanced/expert\n"
            )
        )

        structured_llm = self.llm.with_structured_output(Gaps)
        response = structured_llm.invoke([message])

        gaps_dicts = []
        for gap in response.skills:
            gap_dict = {
                "skill_name": gap.skill_name.strip().lower(),
                "current_level": gap.current_level.strip().lower(),
                "required_level": gap.required_level.strip().lower(),
            }
            gaps_dicts.append(gap_dict)
        return gaps_dicts

    def generate_report(
        self,
        user_profile: Dict[str, Any],
        job_details: Dict[str, str],
        skill_requirements: List[Dict[str, str]],
        skill_gaps: List[Dict[str, str]],
    ) -> Dict:
        return {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "user_name": user_profile["user_name"],
            "user_skills": user_profile["skills"],
            "user_profile_summary": user_profile["profile_summary"],
            "job_role": job_details["job_role"],
            "job_location": job_details["job_location"],
            "company_name": job_details["company_name"],
            "description_summary": job_details["description_summary"],
            "required_skills": skill_requirements,
            "skill_gaps": skill_gaps,
        }


    def analyze(self, job_description, user_cv):
        """User inputs job description and cv"""
        user_profile = self.analyze_cv(user_cv)
        job_role, cleaned_job_description = self.clean_job_description(
            raw_text=job_description
        )
        job_details = self.get_job_details(
            job_role=job_role, job_description=cleaned_job_description
        )
        required_skills = self.get_requirements(cleaned_job_description)
        skill_gaps = self.get_gaps(
            user_profile=user_profile, required_skills=required_skills
        )
        report = self.generate_report(
            user_profile=user_profile,
            job_details=job_details,
            skill_requirements=required_skills,
            skill_gaps=skill_gaps,
        )

        return report
